import os.path
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog
from PyQt5.QtGui import QStandardItemModel, QStandardItem
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ui.top import *
from final_cal.ExcelProcessor import ExcelProcessor, TwoExcelProcess
from final_cal.FilterStrategy import *
from final_cal.ReportGenerator import *
import openpyxl
import pandas as pd
from fuzzywuzzy import fuzz
import win32com.client as win32
import numpy as np
from price.moudle.automator import Automator
from config import Config
from pathlib import Path
import time
import sys
import pyautogui
import win32com.client as win32
from PyQt5.QtCore import Qt, QTimer
from docx import Document
from docx.shared import Cm
import ctypes
import re


class MainWindow(QMainWindow, Ui_MainWindow):
    has_run = False

    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.model = None
        self.setupUi(self)
        self.folderName = ""
        self.keyword = "建筑工程|清单带子目"
        self.directory_paths = []
        self.filter_col = self.FilterColNameLineEdit.text().split(" ")
        self.sum_col = self.SumColNameLineEdit.text()
        self.header = 0

        self.control_and_contract_correspondence_dict = {
            "B1-建筑工程": "管控表1",
            "B2-建筑工程": "管控表2",
            "B3-建筑工程": "管控表3",
            "B4-建筑工程": "管控表4",
            "B5-建筑工程": "管控表5",
            "B6-建筑工程": "管控表6",
            "B7-建筑工程": "管控表7",
            "B8-建筑工程": "管控表8",
            "B9-建筑工程": "管控表9",
            "D1a-建筑工程": "管控表11",
            "D2-建筑工程": "管控表12",
            "D1-建筑工程": "管控表13",
            "D3-建筑工程": "管控表14",
            "B车库-建筑工程": "管控表15",
            "B车库人防-建筑工程": "管控表10",
            "D4-建筑工程": "管控表16",
            "D4a-建筑工程": "管控表17",
            "D4b-建筑工程": "管控表18",
            "D5-建筑工程": "管控表19",
            "D6-建筑工程": "管控表20",
            "D7-建筑工程": "管控表21",
            "D8-建筑工程": "管控表22",
            "D9-建筑工程": "管控表23",
            "D10-建筑工程": "管控表24",
            "D11-建筑工程": "管控表25",
            "D12-建筑工程": "管控表26",
            "D13-建筑工程": "管控表27",
            "D车库-建筑工程": "管控表28",
            "D车库人防-建筑工程": "管控表29",
            "B3a-建筑工程": "管控表34",
        }
        self.word_folder_name = ''
        self.word_pic_paths = []


        self.SelectNeedDataAnalysisDirButton.clicked.connect(self.openFileDialog)
        self.ReruleLineEdit.editingFinished.connect(self.input_re_keyword)
        self.FilterColNameLineEdit.editingFinished.connect(self.input_filter_col)
        self.SumColNameLineEdit.editingFinished.connect(self.input_sum_col)
        self.HeaderSpin.valueChanged.connect(self.input_header)
        self.fianlPushButton.clicked.connect(self.data_main)
        self.five_measure_ledger_button.clicked.connect(self.five_measure_ledger)
        self.five_measure_ledger_button.clicked.connect(self.five_measure_ledger)
        self.testButton.clicked.connect(self.directly_process_the_five_quantity_ledger)
        self.dataMarkButton.clicked.connect(self.data_mark)
        # 计价内容
        self.export_excel_task_imgs_file_names = None
        self.base_path = Config.get_img_base_path()
        self.tabWidget.currentChanged.connect(self.add_radio_button)
        self.add_radio_button()
        self.pushButton_7.clicked.connect(self.function)
        self.pushButton_2.clicked.connect(self.function_add_steps)
        self.pushButton_5.clicked.connect(self.function_delete_steps)
        self.pushButton_6.clicked.connect(self.function_insert_steps)
        self.pushButton_10.clicked.connect(self.clear_function_input)
        self.pushButton_8.clicked.connect(self.clear_delete_input)
        self.pushButton_9.clicked.connect(self.clear_insert_input)

        # 工具类
        # 截图
        self.pushButton.clicked.connect(self.screenshot)

        # 初始化 QGraphicsView 和 QGraphicsScene
        self.scene = QtWidgets.QGraphicsScene(self)

        # 图片显示
        self.graphicsView.setScene(self.scene)

        # 设置 QGraphicsView 的渲染提示
        self.graphicsView.setRenderHint(QtGui.QPainter.Antialiasing)
        self.graphicsView.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        self.graphicsView.setFocusPolicy(QtCore.Qt.StrongFocus)

        # excel筛选命令
        self.pushButton_4.clicked.connect(self.filter_excel)

        self.toolButton_2.clicked.connect(self.toggle_hide)  # 按钮点击事件连接到 toggle_hide 方法

        self.is_hidden = False  # 状态标记，记录窗口是否隐藏

        # word工具内容
        self.pushButton_3.clicked.connect(self.handle_word_pic)
        self.pushButton_12.clicked.connect(self.get_doc)
        self.pushButton_11.clicked.connect(self.get_pic)
        self.word_folder_name = ''
        self.word_pic_paths = []
        self.doc_file_list = None
        self.pic_file_list = None
        self.pic_folder_name = ''

    def keyPressEvent(self, event):
        try:
            # 处理 Ctrl + V (粘贴)
            if event.modifiers() == QtCore.Qt.ControlModifier and event.key() == QtCore.Qt.Key_V:
                clipboard = QtWidgets.QApplication.clipboard()
                mime_data = clipboard.mimeData()

                if mime_data.hasImage():
                    image = mime_data.imageData()
                    pixmap = QtGui.QPixmap(image)

                    # 获取鼠标位置并转换为场景坐标
                    if self.graphicsView.hasFocus():
                        mouse_pos = self.graphicsView.mapFromGlobal(QtGui.QCursor.pos())
                        scene_pos = self.graphicsView.mapToScene(mouse_pos)
                        scene = self.scene
                    else:
                        return

                    # 将图片添加到场景
                    pixmap_item = scene.addPixmap(pixmap)
                    pixmap_item.setFlag(pixmap_item.ItemIsMovable)
                    pixmap_item.setFlag(pixmap_item.ItemIsSelectable)
                    pixmap_item.setPos(scene_pos)

                    event.accept()
                else:
                    super().keyPressEvent(event)

            # 处理 Ctrl + X (剪切)
            elif event.modifiers() == QtCore.Qt.ControlModifier and event.key() == QtCore.Qt.Key_X:
                # 判断哪个 QGraphicsView 获取了焦点
                if self.graphicsView.hasFocus():
                    scene = self.scene
                else:
                    return
                # 获取选中的图形项
                selected_items = scene.selectedItems()
                if selected_items:
                    # 假设只处理一个选中的项
                    pixmap_item = selected_items[0]
                    if isinstance(pixmap_item, QtWidgets.QGraphicsPixmapItem):
                        # 将选中的项复制到剪贴板
                        pixmap = pixmap_item.pixmap()
                        clipboard = QtWidgets.QApplication.clipboard()
                        clipboard.setPixmap(pixmap)

                        # 从场景中删除选中的项
                        scene.removeItem(pixmap_item)

                        event.accept()
                    else:
                        super().keyPressEvent(event)
                else:
                    super().keyPressEvent(event)

        except Exception as e:
            print(f"Error: {e}")

    def toggle_hide(self):
        """ 切换窗口是否隐藏到屏幕右侧 """
        if self.is_hidden:
            self.showNormal()  # 恢复窗口显示
            # 调整窗口的位置，确保它恢复到屏幕内
            screen_geometry = QApplication.primaryScreen().geometry()  # 获取屏幕的大小
            self.move(screen_geometry.width() - self.width(), self.y())  # 确保窗口恢复在屏幕内
        else:
            self.hide_to_right()  # 隐藏窗口到屏幕右侧
        self.is_hidden = not self.is_hidden  # 切换状态

    def hide_to_right(self):
        """ 将窗口隐藏到屏幕的右侧外部 """
        screen_geometry = QApplication.primaryScreen().geometry()  # 获取屏幕的大小
        screen_width = screen_geometry.width()
        screen_height = screen_geometry.height()

        # 设置窗口位置为屏幕右侧外部
        self.move(screen_width - 30, 0)  # 将窗口移到屏幕的右侧外面

    def closeEvent(self, event):
        """ 在窗口关闭时，保持窗口在右侧外部 """
        if self.is_hidden:
            self.hide_to_right()
            event.ignore()  # 忽略关闭事件，防止程序退出
        else:
            event.accept()  # 正常关闭窗口

    def openFileDialog(self):
        # 打开文件夹选择对话框
        self.folderName = QFileDialog.getExistingDirectory(self, "选择文件夹", "")
        if self.folderName:
            # print(f'Selected folder: {self.folderName}')
            self.SelectPathLineEdit.setText(self.folderName)
            self.directory_paths.append(self.folderName.replace("/", "\\"))
            self.show_columns(self.directory_paths, self.keyword, self.header)

    def show_columns(self, directory_paths, keyword, header):
        try:
            file_manager = FileManager(directory_paths, keyword)
            excel_files = file_manager.get_excel_files()
            data = file_manager.read_excel(excel_files[0], header=header)
            self.model = DataFrameModel(data)
            self.TableColName.setModel(self.model)
        except Exception as e:
            print(e)

    def input_re_keyword(self):
        self.keyword = self.ReruleLineEdit.text()

    def input_filter_col(self):
        self.filter_col = self.FilterColNameLineEdit.text().split(" ")

    def input_sum_col(self):
        self.sum_col = self.SumColNameLineEdit.text()

    def input_header(self):
        self.header = int(self.HeaderSpin.text())
        if self.directory_paths:
            self.show_columns(self.directory_paths, self.keyword, self.header)

    def data_main(self):
        output_path = os.path.join(self.directory_paths[0], 'out.xlsx')
        keyword = self.keyword
        report_generator = SubReport()
        filter_strategy = None
        if not self.checkBox.isChecked() and not self.checkBox_7.isChecked():
            # filter_strategy = GeneralFilter([(4, "砖基础"), (3, "001")])
            filter_strategy = NameProjectFeatureUnitSameFilter(self.filter_col, self.sum_col)

        elif self.checkBox.isChecked():
            # self.SumColNameLineEdit.clear()
            file_manager = FileManager(self.directory_paths, self.keyword)
            excel_files = file_manager.get_excel_files()
            if len(excel_files) > 1:
                self.compare_two_excel_sheets(excel_files[0], excel_files[1])

        elif self.checkBox_7.isChecked():
            filter_strategy = NameProjectFeatureUnitSameFilterList(self.filter_col, self.sum_col)

        processor = ExcelProcessor(filter_strategy, report_generator)
        summed_data = processor.process(self.directory_paths, output_path, keyword, self.header)
        self.model = DataFrameModel(summed_data)
        self.DataTableView.setModel(self.model)

    def compare_two_excel_sheets(self, df1_path, df2_path):
        df1_file_name = df1_path.split('\\')[-1].split('.')[0]
        df2_file_name = df2_path.split('\\')[-1].split('.')[0]
        # 读取两个 Excel 文件
        df1 = pd.read_excel(df1_path)  # 替换成实际文件路径
        df2 = pd.read_excel(df2_path)  # 替换成实际文件路径

        # 去除列名和数据中的多余空格
        df1.columns = df1.columns.str.strip()
        df2.columns = df2.columns.str.strip()

        df1 = df1.apply(lambda x: x.strip() if isinstance(x, str) else x)
        df2 = df2.apply(lambda x: x.strip() if isinstance(x, str) else x)

        # 合并数据：内连接（匹配行）
        df_matched = pd.merge(df1, df2, how='inner', left_on=self.filter_col, right_on=self.filter_col)

        # 找出df1有而df2没有的行（左连接）
        df1_only = pd.merge(df1, df2, how='left', left_on=self.filter_col, right_on=self.filter_col,
                            indicator=True)
        df1_only = df1_only[df1_only['_merge'] == 'left_only'].drop(columns=['_merge'])

        # 找出df2有而df1没有的行（右连接）
        df2_only = pd.merge(df1, df2, how='right', left_on=self.filter_col, right_on=self.filter_col,
                            indicator=True)
        df2_only = df2_only[df2_only['_merge'] == 'right_only'].drop(columns=['_merge'])

        # 保存结果
        df_matched.to_excel(os.path.join(self.folderName, f"{df1_file_name}{df2_file_name}都有.xlsx"), index=False)
        df1_only.to_excel(os.path.join(self.folderName, f"{df1_file_name}有{df2_file_name}没有.xlsx"), index=False)
        df2_only.to_excel(os.path.join(self.folderName, f"{df2_file_name}有{df1_file_name}没有.xlsx"), index=False)

        print("文件已保存！")

    def five_measure_ledger(self):
        for checklist, control in self.control_and_contract_correspondence_dict.items():
            self.five_measure_ledger_circulation(checklist, control)
            print(checklist + '处理完成！')

    def five_measure_ledger_circulation(self, checklist, control):
        # 打开已有的工作簿
        lou = checklist.split("-")[0]
        path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\5、结算\6、总预算\只是清单\辉县市灾后重建项目城南片区安置区建设项目(BD地块）（审核）\{lou}\{checklist}\{checklist}.xlsx'
        workbook = openpyxl.load_workbook(path)

        # 选择工作表
        sheet = workbook['表-08 分部分项工程和单价措施项目清单与计价表']

        # 在 F 列后插入三列
        # openpyxl 的列索引从 1 开始，F 列是第 6 列，所以我们在第 7 列（G 列）之前插入
        sheet.insert_cols(7, amount=3)
        sheet.delete_cols(11, 2)
        # 保存工作簿
        workbook.save(path)

        # 读取 Excel 文件
        excel_file = path
        df = pd.read_excel(excel_file, sheet_name='表-08 分部分项工程和单价措施项目清单与计价表')

        # 删除包含“合计”和“本页小计”的行
        df = df[~df.iloc[:, 0].astype(str).str.contains('合计|本页小计|顺序号', na=False)]
        df = df[~df.iloc[:, 2].astype(str).str.contains('分部小计|单价措施', na=False)]
        df = df[~df.iloc[:, 0].astype(str).str.contains(r'合*计', na=False)]
        df = df[~df.iloc[:, 9].astype(str).str.contains('综合单价', na=False)]
        df = df[~df.iloc[:, 9].astype(str).str.contains('综合单价', na=False)]

        # 删除“名称”和“编码”均为空的行
        df = df.dropna(subset=['项目名称', '项目编码'], how='all')

        # 保存处理后的数据到新的 Excel 文件
        out_path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\5、结算\6、总预算\只是清单\辉县市灾后重建项目城南片区安置区建设项目(BD地块）（审核）\{lou}\{checklist}\out.xlsx'
        df.to_excel(out_path, index=False, sheet_name='表-08 分部分项工程和单价措施项目清单与计价表')

        # 将管控表列名“项目名称及细目细目特征”列拆分为项目名称、细目细目特征进行存储
        control_tables_path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\2024.1.22-财评对量\2024.6.4\总量\五量台账\管控表\{control}.xlsx'
        control_table = pd.read_excel(control_tables_path, sheet_name=f'{control}', header=2)

        # 去掉列名不必要字符
        # 清理列名的空字符、换行符、制表符等
        control_table.columns = (control_table.columns
                                 .str.strip()  # 去掉前后空格
                                 .str.replace('\n', '', regex=False)  # 去掉换行符
                                 .str.replace('\t', '', regex=False))  # 去掉制表符

        # 获取第二个列名（索引为1）
        second_column = control_table.columns[1]

        # 替换第二个列名中的 "_x000D_" 为 ""
        updated_second_column = second_column.replace('_x000D_', '')

        # 更新 DataFrame 的列名
        columns = list(control_table.columns)
        columns[1] = updated_second_column
        control_table.columns = columns

        # 获取第三个列名（索引为2）
        second_column = control_table.columns[2]

        # 替换第二个列名中的 "_x000D_" 为 ""
        updated_second_column = second_column.replace('_x000D_', '')

        # 更新 DataFrame 的列名
        columns = list(control_table.columns)
        columns[2] = updated_second_column
        control_table.columns = columns

        # 遍历预算文件将上面处理好的预算文件（out.xlsx）与管控表进行对比
        # 如果管控表项目名称（project_name）和预算的项目名称相等且管控表特征描述（features_description）和预算的特征描述相似度达90%进行下一步
        budget_file = pd.read_excel(out_path, sheet_name='表-08 分部分项工程和单价措施项目清单与计价表', header=0)

        new_budget = control_table.copy()

        # 保留前两行数据
        new_budget = new_budget.iloc[:2]

        other_budget = budget_file.copy()
        other_budget.loc[:] = np.nan

        count = 2

        other_count = 0

        for budget_idx, budget_value in budget_file['项目特征描述'].items():
            # 添加列名“工程量清单编号”不为空的条件
            temp_list = str(budget_value).split('\n')
            budget_value_clean = ''.join([item for item in temp_list if '其他:' not in item]).replace(" ", "").replace(
                "\n", "").replace("\r", "").strip()
            for control_idx, control_value in control_table['项目名称及细目细目特征'].astype(str).items():
                if pd.notna(control_table.loc[control_idx, '工程量清单编号']):
                    temp_split = control_value.split('1.')
                    project_name = temp_split[0]
                    features_description = '1.' + temp_split[-1]
                    control_value_clean = str(features_description).replace(" ", "").replace("\n", "").replace("\r",
                                                                                                               "").strip()
                    similarity = fuzz.ratio(budget_value_clean, control_value_clean)
                    similarity_project_name = fuzz.ratio(budget_file.at[budget_idx, '项目名称'], project_name)
                    if similarity_project_name >= 90 and similarity >= 100:
                        # 在预算该行下插入x行，x行由管控表对应行到下一个项目编码有值的行数决定
                        # 计算插入行数
                        # 确保数据的行数大于 1
                        # 找到列中非空值的所有索引
                        non_null_indices = control_table.index[control_table['工程量清单编号'].notnull()].tolist()

                        # 确保已知的非空索引在非空值列表中
                        if control_idx not in non_null_indices:
                            raise ValueError("已知的非空值索引不在列中。")

                        # 找到该索引的位置
                        try:
                            idx_position = non_null_indices.index(control_idx)
                            # 获取下一个非空索引
                            next_non_null_index = non_null_indices[idx_position + 1] if idx_position < len(
                                non_null_indices) - 1 else None
                        except IndexError:
                            next_non_null_index = None

                        new_budget.at[count, '章节号'] = budget_file.at[budget_idx, '序号']
                        new_budget.at[count, '工程量清单编号'] = budget_file.at[budget_idx, '项目编码']
                        new_budget.at[count, '项目名称及细目细目特征'] = budget_file.at[budget_idx, '项目名称'] + budget_file.at[
                            budget_idx, '项目特征描述']
                        new_budget.at[count, '单位'] = budget_file.at[budget_idx, '项目名称'] + budget_file.at[
                            budget_idx, '计量单位']
                        new_budget.at[count, '总控工程数量'] = budget_file.at[budget_idx, '工程量']
                        new_budget.at[count, 'Unnamed: 6'] = budget_file.at[budget_idx, '工程量']
                        new_budget.at[count, 'Unnamed: 8'] = budget_file.at[budget_idx, '工程量']
                        new_budget.at[count, '施工合同单价'] = budget_file.at[budget_idx, '金 额(元)']
                        new_budget.at[count, '统计参数设置'] = control_table.at[control_idx, '统计参数设置']
                        new_budget.at[count, '过程数量管控（开累）'] = budget_file.at[budget_idx, '工程量']
                        for add_num in range(1, next_non_null_index - control_idx + 1):
                            count += 1
                            new_budget.at[count, '分包清单序号'] = control_table.at[control_idx + add_num, '分包清单序号']
                            new_budget.at[count, '项目名称及细目细目特征'] = control_table.at[control_idx + add_num, '项目名称及细目细目特征']
                            new_budget.at[count, '单位'] = control_table.at[control_idx + add_num, '单位']
                            new_budget.at[count, 'Unnamed: 6'] = new_budget.at[count, 'Unnamed: 6']
                            new_budget.at[count, 'Unnamed: 8'] = new_budget.at[count, 'Unnamed: 8']
                            new_budget.at[count, '分包合同单价'] = control_table.at[control_idx + add_num, '分包合同单价']
                            new_budget.at[count, '统计参数设置'] = control_table.at[control_idx + add_num, '统计参数设置']
                            new_budget.at[count, 'Unnamed: 14'] = control_table.at[control_idx + add_num, 'Unnamed: 14']
                            new_budget.at[count, 'Unnamed: 15'] = control_table.at[control_idx + add_num, 'Unnamed: 15']
                            new_budget.at[count, '主材1应耗量明细'] = control_table.at[control_idx + add_num, '主材1应耗量明细']
                            new_budget.at[count, 'Unnamed: 20'] = control_table.at[control_idx + add_num, 'Unnamed: 20']
                            new_budget.at[count, 'Unnamed: 21'] = control_table.at[control_idx + add_num, 'Unnamed: 21']
                            new_budget.at[count, '主材2应耗量明细'] = control_table.at[control_idx + add_num, '主材2应耗量明细']
                            new_budget.at[count, 'Unnamed: 23'] = control_table.at[control_idx + add_num, 'Unnamed: 23']
                            new_budget.at[count, 'Unnamed: 24'] = control_table.at[control_idx + add_num, 'Unnamed: 24']

                    elif similarity_project_name >= 90 and (89 < similarity < 100):
                        other_count += 1
                        # print(budget_idx, other_count)
                        if other_count < len(budget_file):
                            other_budget.iloc[other_count] = budget_file.iloc[budget_idx]

        # 去除重复行
        new_budget = new_budget.drop_duplicates()
        new_budget_path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\5、结算\6、总预算\处理后\管控表\{lou}-{control}.xlsx'
        new_budget.to_excel(new_budget_path, index=False, sheet_name=f'{control}')

        other_budget = other_budget.drop_duplicates()
        other_budget_path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\5、结算\6、总预算\处理后\管控表\other-{lou}-{control}.xlsx'
        other_budget.to_excel(other_budget_path, index=False, sheet_name=f'{control}')

    def directly_process_the_five_quantity_ledger(self):
        excel_dict = {
            "B1-建筑工程": "管控表1",
            "B2-建筑工程": "管控表2",
            "B3-建筑工程": "管控表3",
            "B4-建筑工程": "管控表4",
            "B5-建筑工程": "管控表5",
            "B6-建筑工程": "管控表6",
            "B7-建筑工程": "管控表7",
            "B8-建筑工程": "管控表8",
            "B9-建筑工程": "管控表9",
            "D1a-建筑工程": "管控表11",
            "D2-建筑工程": "管控表12",
            "D1-建筑工程": "管控表13",
            "D3-建筑工程": "管控表14",
            "B车库-建筑工程": "管控表15",
            "D4-建筑工程": "管控表16",
            "D4a-建筑工程": "管控表17",
            "D4b-建筑工程": "管控表18",
            "D5-建筑工程": "管控表19",
            "D6-建筑工程": "管控表20",
            "D7-建筑工程": "管控表21",
            "D8-建筑工程": "管控表22",
            "D9-建筑工程": "管控表23",
            "D10-建筑工程": "管控表24",
            "D11-建筑工程": "管控表25",
            "D12-建筑工程": "管控表26",
            "D13-建筑工程": "管控表27",
            "D车库-建筑工程": "管控表28",
            "D车库人防-建筑工程": "管控表29",
            "B3a-建筑工程": "管控表34",
        }

        path_control_tables = r'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\2024.1.22-财评对量\2024.6.4\总量\五量台账\管控表'
        path_budget = r'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\2024.1.22-财评对量\2024.6.4\总量\五量台账\预算'

        for k, v in excel_dict.items():
            # 读取两个 Excel 文件
            file1 = path_control_tables + fr'\{v}.xlsx'
            file2 = path_budget + fr'\{k}\表-09 综合单价分析表(清单带子目).xlsx'

            df1 = pd.read_excel(file1, sheet_name=f'{v}', header=2)
            df2 = pd.read_excel(file2, sheet_name='表-09 综合单价分析表(清单带子目)', header=0)

            # 初始化结果DataFrame
            df1_modified = df1.copy()

            # 比较 example1 的 D 列和 example2 的 D 列
            for i, value1 in df1['项目名称及细目细目特征'].items():
                value1_clean = str(value1).replace(" ", "").replace("\n", "").replace("\r", "").strip()
                for j, value2 in df2['项目特征描述'].items():
                    value2_clean = str(value2).replace(" ", "").replace("\n", "").replace("\r", "").strip()
                    value2_name_features = str(df2.loc[j, '名称']) + value2_clean
                    similarity = fuzz.ratio(value1_clean, value2_name_features)
                    if similarity >= 80:
                        df1_modified.at[i, '单位'] = df2.at[j, '计量单位']
                        df1_modified.at[i, '总控工程数量'] = df2.at[j, '工程量']
                        df1_modified.at[i, '过程数量管控（开累）'] = df2.at[j, '工程量']
                        df1_modified.at[i, 'Unnamed: 6'] = df2.at[j, '工程量']
                        df1_modified.at[i, 'Unnamed: 8'] = df2.at[j, '工程量']
                        df1_modified.at[i, '施工合同单价'] = df2.at[j, '金额（元）']
                        break  # 找到一个匹配就跳出内循环，处理下一个值
            out_path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\2024.1.22-财评对量\2024.6.4\总量\五量台账\导出\{v}.xlsx'
            # 保存处理后的数据到新的 Excel 文件
            df1_modified.to_excel(out_path, index=False, sheet_name=f'{v}')

    def data_mark(self):
        for _, v in self.control_and_contract_correspondence_dict.items():
            file_path = fr'D:\Work_Content\新乡灾后重建安置房项目\44、工程数量台账化管理\2024.1.22-财评对量\2024.6.4\总量\五量台账\导出\{v}.xlsx'
            sheet_name = v
            column_letter = "F"
            # 启动 Excel 应用程序
            excel = win32.Dispatch('Excel.Application')
            excel.Visible = False  # 不显示 Excel 窗口
            wb = excel.Workbooks.Open(file_path)
            ws = wb.Sheets(sheet_name)

            # 获取最后一行
            xlUp = -4162
            last_row = ws.Cells(ws.Rows.Count, column_letter).End(xlUp).Row

            # 创建一个字典来存储每个值的出现次数
            value_counts = {}

            # 遍历列中的每个单元格并统计出现次数
            for row in range(2, last_row + 1):  # 从第2行开始，假设第1行为标题
                cell_value = ws.Cells(row, column_letter).Value
                if cell_value is not None:
                    if cell_value in value_counts:
                        value_counts[cell_value] += 1
                    else:
                        value_counts[cell_value] = 1

            # 遍历列中的每个单元格并应用格式
            for row in range(2, last_row + 1):
                cell_value = ws.Cells(row, column_letter).Value
                if cell_value is not None and value_counts[cell_value] >= 2:
                    ws.Cells(row, column_letter).Interior.Color = 0x00FF00  # 绿色

            # 保存并关闭工作簿
            wb.Save()
            wb.Close(False)
            excel.Quit()

    # 计价内容
    # 执行
    def run_task(self, c_task_name):
        automator = Automator()
        automator.execute_task(c_task_name)

    def task(self, c_task_img_path):
        # print(c_task_img_path)
        export_excel_task_path = Path(self.base_path) / c_task_img_path
        # 获取目录下所有文件的文件名（只取文件）
        self.export_excel_task_imgs_file_names = [f for f in export_excel_task_path.iterdir() if f.is_file()]
        # print([f.name for f in export_excel_task_imgs_file_names])

        # 根据序号顺序执行任务
        for img in sorted(self.export_excel_task_imgs_file_names, key=self.extract_number):  # 按文件名.排序
            Automator.task_img_path = img.resolve()
            print(Automator.task_img_path)
            task_name = img.stem.split('-')[-1]  # 只取文件名，不含扩展名
            # print(task_name)
            try:
                self.run_task(task_name)
            except Exception as e:
                print(f"任务 {task_name} 执行失败: {e}")

    def extract_number(self, s):
        """ 提取文件名中的数字部分并返回数字值，用于排序 """
        # 确保提取的是文件名的字符串部分
        file_name = str(s.name)  # 转换为字符串（文件名）
        match = re.match(r"(\d+)", file_name)  # 提取文件名开头的数字部分
        # 如果没有数字，返回一个很大的数字，让它排到最后
        return int(match.group(1)) if match else float('inf')

    def add_radio_button(self):
        self.textEdit.clear()
        for radio_button_obj in self.widget_19.findChildren(QtWidgets.QRadioButton):
            # 删除 QRadioButton
            radio_button_obj.deleteLater()
        task_menu = self.get_tasks_menu()
        # print(task_menu)
        for task_path in task_menu:
            radio_button_name = f'radio_button_{task_path.stem}'
            self.radioButton = QtWidgets.QRadioButton(self.widget_19)
            self.radioButton.setObjectName(radio_button_name)
            self.horizontalLayout_14.addWidget(self.radioButton)
            _translate = QtCore.QCoreApplication.translate
            self.radioButton.setText(_translate("MainWindow", task_path.stem))
            self.radioButton.clicked.connect(self.show_already_task_steps)

    def get_already_task_steps(self) -> list:
        task_name_list = []
        radio_button = self.get_selected_radio_button()
        file_path = Path(self.base_path) / radio_button.text()
        file = [f for f in file_path.iterdir() if f.is_file()]
        for img in sorted(file, key=lambda x: x.stem):  # 按文件名排序
            task_name = img.stem  # 只取文件名，不含扩展名
            task_name_list.append(task_name)
        return task_name_list

    def show_already_task_steps(self):
        self.textEdit.clear()
        task_name_list = self.get_already_task_steps()
        for task_name in task_name_list:
            self.textEdit.append(task_name)

    def get_tasks_menu(self):
        tasks_menu = [menu for menu in Path(self.base_path).iterdir() if menu.is_dir()]

        return tasks_menu

    def function(self):
        # # c_task_name为当前视口被选中的raidobutton
        # task_name_path = Path(self.base_path) / c_task_name
        # if not os.path.exists(task_name_path):
        #     os.mkdir(task_name_path)
        radio_button = self.get_selected_radio_button()
        self.task(radio_button.text())

    def get_selected_radio_button(self):
        for radio_button in self.widget_19.findChildren(QtWidgets.QRadioButton):
            # print(radio_button)
            if radio_button.isChecked():  # 检查是否被选中
                return radio_button

    def function_add_steps(self):
        try:
            if not self.checkBox_3.isChecked():
                # 增加自动编顺序
                radio_button = self.get_selected_radio_button()
                task_name_list = self.get_already_task_steps()
                # print(task_name_list)
                num = len(task_name_list)
                if self.lineEdit.text() != '':
                    step = f'{num}-{self.lineEdit.text()}.png'
                    # print(step)
                    step_path = Path(self.base_path) / radio_button.text() / step
                    if len(self.scene.items()) != 0:
                        for item in self.scene.items():
                            # print(".......")
                            # print(step)
                            if isinstance(item, QtWidgets.QGraphicsPixmapItem):
                                pixmap = item.pixmap()
                                if not pixmap.isNull():
                                    # 保存图片到指定文件
                                    pixmap.save(str(step_path))
                            else:
                                # 其他内容
                                pass
                    else:
                        with open(str(step_path), 'w') as fp:
                            pass
                else:
                    # 提示输入
                    pass
            else:
                # 创建任务目录
                new_task_dir = os.path.join(self.base_path, self.lineEdit.text())
                os.mkdir(new_task_dir)
        except Exception as e:
            print(e)

    def function_insert_steps(self):
        try:
            # 插入索引在idx位置
            idx = int(self.spinBox_3.text())
            print(idx)
            radio_button = self.get_selected_radio_button()
            if self.lineEdit_2.text() != '':
                step_name = f'{idx}-{self.lineEdit_2.text()}.png'
                for item in self.scene.items():
                    step_path = Path(self.base_path) / radio_button.text() / step_name
                    # print(step)
                    if isinstance(item, QtWidgets.QGraphicsPixmapItem):
                        pixmap = item.pixmap()
                        if not pixmap.isNull():
                            # 保存图片到指定文件
                            pixmap.save(str(step_path))
                    else:
                        with open(str(step_path), 'w') as fp:
                            pass
                # 重名插入后索引，索引加一
                already_steps_list = self.get_already_task_steps()[idx:]  # 获取原来文件名列表
                need_rename_list = [x.split('-')[-1] for x in already_steps_list]  # 需改名部位
                # print(need_rename_list)
                old_name_path = [Path(self.base_path) / radio_button.text() / f'{path}.png' for i, path in
                                 enumerate(already_steps_list)]
                # if idx == 0:
                #     idx = -1
                new_name_path_list = [Path(self.base_path) / radio_button.text() / f'{idx + i}-{path}.png' for
                                      i, path in enumerate(need_rename_list)]
                # print(old_name_path)
                # print(new_name_path_list)
                # print(*zip(old_name_path, new_name_path_list))
                for o, n in zip(old_name_path, new_name_path_list):
                    os.rename(str(o), str(n))
            else:
                # 提示输入
                pass
        except Exception as e:
            print(e)

    def function_delete_steps(self):
        try:
            radio_button = self.get_selected_radio_button()
            task_name_list = self.get_already_task_steps()
            # 去序号
            task_name_no_num_list = [name.split("-")[-1] for name in task_name_list]
            if not self.checkBox_2.isChecked() and not self.checkBox_4.isChecked():
                if self.lineEdit_3.text() in task_name_no_num_list:
                    current_step_name = [f for f in task_name_list if self.lineEdit_3.text() in f][0]
                    current_step_name_fin = f'{current_step_name}.png'
                    # 获取当前步骤的路径
                    current_task = radio_button.text()
                    current_step_name = self.lineEdit_3.text()
                    current_step_path = os.path.join(os.path.join(self.base_path, current_task), current_step_name_fin)
                    os.remove(current_step_path)

                    task_name_list = self.get_already_task_steps()
                    # 更新已有序号 重命名
                    for i, old_name in enumerate(task_name_list):
                        print(old_name)
                        old_path = os.path.join(os.path.join(self.base_path, current_task), f'{old_name}.png')
                        new_name = os.path.join(os.path.join(self.base_path, current_task),
                                                f'{i}-{old_name.split("-")[-1]}.png')
                        if os.path.exists(old_path):
                            os.rename(old_path, new_name)
                            print(f"文件已重命名为 {new_name}")
                        else:
                            print(f"文件 {old_name} 不存在")

                else:
                    print("steps不存在")
            elif self.checkBox_2.isChecked():
                # 确保目录存在
                radio_button = self.get_selected_radio_button()
                directory = Path(self.base_path) / radio_button.text()
                if not os.path.exists(directory):
                    print(f"目录 {directory} 不存在")
                    return

                # 遍历目录中的所有文件
                for filename in os.listdir(directory):
                    file_path = os.path.join(directory, filename)

                    # 检查是否是文件而不是目录
                    if os.path.isfile(file_path):
                        try:
                            os.remove(file_path)
                            print(f"已删除文件: {file_path}")
                        except Exception as e:
                            print(f"删除文件 {file_path} 失败: {e}")
                    else:
                        print(f"{file_path} 不是文件，跳过")
            elif self.checkBox_4.isChecked():
                # 删除任务
                del_task_path = os.path.join(self.base_path, radio_button.text())
                os.remove(del_task_path)

        except Exception as e:
            print(e)

    def clear_function_input(self):
        self.lineEdit.clear()

    def clear_delete_input(self):
        self.lineEdit_3.clear()

    def clear_insert_input(self):
        self.lineEdit_2.clear()

    # 截图
    def screenshot(self):
        pyautogui.hotkey('win', 'shift', 's')

    def filter_excel(self):
        # 打开Excel应用
        excel = win32.Dispatch('Excel.Application')

        # 如果不想显示Excel界面，可以设置为False
        excel.Visible = True

        # 打开现有的工作簿
        wb = excel.ActiveWorkbook

        sheet = wb.ActiveSheet  # 选择当前活动工作表

        # 设置筛选范围（例如A1到D100）
        data_range = sheet.Range('A1:Z100')

        # 启用筛选
        data_range.AutoFilter()

        # 检查是否有筛选器
        if sheet.AutoFilterMode:
            # 如果有筛选器，取消筛选
            sheet.AutoFilter.ShowAllData()

    def get_doc(self):
        try:
            # 打开文件选择对话框，限制只能选择 .doc 或 .docx 文件
            options = QFileDialog.Options()
            file_paths, _ = QFileDialog.getOpenFileNames(
                None,
                "选择 Word 文件",
                "",
                "Word Files (*.docx *.doc)",
                options=options
            )
            self.doc_file_list = file_paths
            self.word_folder_name = Path(file_paths[0]).parent
            if self.word_folder_name:
                # print(f'Selected folder: {self.folderName}')
                self.lineEdit_5.setText(str(self.word_folder_name))
        except Exception as e:
            print(e)

    def get_pic(self):
        try:
            # 打开文件选择对话框，限制只能选择 .doc 或 .docx 文件
            options = QFileDialog.Options()
            file_paths, _ = QFileDialog.getOpenFileNames(
                None,
                "选择 图片 文件",
                "",
                "Word Files (*.jpg *.jpeg *.png *.webp)",
                options=options
            )
            self.pic_file_list = file_paths
            self.pic_folder_name = Path(file_paths[0]).parent
            if self.pic_folder_name:
                # print(f'Selected folder: {self.folderName}')
                self.lineEdit_4.setText(str(self.pic_folder_name))
        except Exception as e:
            print(e)

    def handle_word_pic(self):
        try:
            doc_list = self.doc_file_list
            pic_list = self.pic_file_list
            if not self.checkBox_6.isChecked():
                # 打开 Word 文档
                doc = Document(doc_list[0])
                # 遍历文档中的每个段落
                for par in doc.paragraphs:
                    # 遍历每个段落中的每个运行（run）
                    for run in par.runs:
                        if run.element.drawing_lst is not None:  # 如果有图片
                            # 设置段落对齐方式为居中
                            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 遍历文档中的所有内嵌图片
                for inline_shape in doc.inline_shapes:
                    # 设置图片大小 (可选)
                    inline_shape.width = Cm(int(self.spinBox.text()))  # 设置图片宽度为5厘米
                    inline_shape.height = Cm(3)  # 设置图片高度为3厘米
                # 保存文档
                doc.save(doc_list[0])
            else:
                # 打开 Word 文档
                doc = Document(doc_list[0])
                for p in pic_list:
                    para = doc.add_paragraph()
                    para.alignment = 1  # 1 表示居中对齐
                    # 插入第一张图片
                    para.add_run().add_picture(p, width=Cm(int(self.spinBox.text())), height=Cm(int(self.spinBox_2.text())))
                # 保存文档
                doc.save(doc_list[0])
        except Exception as e:
            print(e)


class DataFrameModel(QStandardItemModel):
    def __init__(self, dataFrame):
        super().__init__()
        self._dataFrame = dataFrame
        self.load_data()

    def load_data(self):
        self.setHorizontalHeaderLabels(self._dataFrame.columns.tolist())
        for row_idx, row in self._dataFrame.iterrows():
            items = [QStandardItem(str(field)) for field in row]
            self.appendRow(items)


def is_admin():
    """检查是否为管理员权限"""
    return ctypes.windll.shell32.IsUserAnAdmin() != 0


def run_as_admin():
    """以管理员权限重新启动当前脚本"""
    script = sys.argv[0]
    params = ' '.join(sys.argv[1:])
    ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, f'"{script}" {params}', None, 1)


if __name__ == '__main__':
    # # 检查是否具有管理员权限
    # if not is_admin():
    #     print("需要管理员权限，正在以管理员权限重新启动...")
    #     run_as_admin()  # 重新以管理员权限运行脚本
    #     sys.exit(0)
    # # # 设置应用程序的背景透明，确保窗口和控件都能显示
    # # app.setAttribute(Qt.AA_UseSoftwareOpenGL)
    app = QApplication(sys.argv)
    myWin = MainWindow()
    myWin.show()
    sys.exit(app.exec_())


